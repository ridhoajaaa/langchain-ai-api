"""
LangChain FastAPI Server
=========================
Full-stack AI server with Chat, Agent, RAG, PDF upload,
conversation memory, and tools.
Run:  uvicorn server:app --reload --host 0.0.0.0 --port 8000
"""

import os
import json
import logging
import uuid
from typing import Any
from contextlib import asynccontextmanager
from pathlib import Path

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from sse_starlette.sse import EventSourceResponse

load_dotenv()

# ── Logging ────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
log = logging.getLogger("langchain-server")

# ── Constants ──────────────────────────────────────────────────
MAX_HISTORY_PER_SESSION = 20  # Keep last N messages for context
UPLOAD_DIR = Path("/tmp/langchain_uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════
#  LLM & Provider Factory
# ═══════════════════════════════════════════════════════════════

def create_llm(provider: str = "ollama", **kwargs) -> Any:
    """Factory to create LLM from provider name."""
    model = kwargs.pop("model", None)

    if provider == "ollama":
        from langchain_ollama import ChatOllama
        default_model = os.getenv("OLLAMA_MODEL", "qwen3-coder-next:cloud")
        log.info("Creating Ollama LLM with model=%s", model or default_model)
        return ChatOllama(
            model=model or default_model,
            temperature=kwargs.get("temperature", 0.7),
            base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
        )

    elif provider == "openai":
        from langchain_openai import ChatOpenAI
        return ChatOpenAI(
            model=model or "gpt-4o-mini",
            temperature=kwargs.get("temperature", 0.7),
            api_key=os.getenv("OPENAI_API_KEY"),
        )

    elif provider == "anthropic":
        from langchain_anthropic import ChatAnthropic
        return ChatAnthropic(
            model=model or "claude-sonnet-4-20250514",
            temperature=kwargs.get("temperature", 0.7),
            api_key=os.getenv("ANTHROPIC_API_KEY"),
        )

    elif provider == "google":
        from langchain_google_genai import ChatGoogleGenerativeAI
        return ChatGoogleGenerativeAI(
            model=model or "gemini-2.0-flash",
            temperature=kwargs.get("temperature", 0.7),
            api_key=os.getenv("GOOGLE_API_KEY"),
        )
    else:
        raise ValueError(f"Unknown provider: {provider}. Use: ollama, openai, anthropic, google")


# ═══════════════════════════════════════════════════════════════
#  Tools
# ═══════════════════════════════════════════════════════════════

from langchain_core.tools import tool


@tool
def calculate(expression: str) -> str:
    """Evaluate a mathematical expression. Input: math expression string."""
    import ast
    import operator as op

    operators = {
        ast.Add: op.add, ast.Sub: op.sub, ast.Mult: op.mul,
        ast.Div: op.truediv, ast.Pow: op.pow, ast.USub: op.neg,
        ast.Mod: op.mod, ast.FloorDiv: op.floordiv,
    }

    def _eval(node):
        if isinstance(node, ast.Constant):
            return node.n if isinstance(node.n, (int, float)) else node.value
        elif isinstance(node, ast.BinOp):
            return operators[type(node.op)](_eval(node.left), _eval(node.right))
        elif isinstance(node, ast.UnaryOp):
            return operators[type(node.op)](_eval(node.operand))
        elif isinstance(node, ast.Expression):
            return _eval(node.body)
        else:
            raise TypeError(f"Unsupported: {type(node).__name__}")

    try:
        tree = ast.parse(expression, mode="eval")
        result = _eval(tree.body)
        return f"Hasil: {result}"
    except Exception as e:
        return f"Error: {e}"


@tool
def get_current_time() -> str:
    """Get the current date and time."""
    from datetime import datetime
    return f"Sekarang: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"


@tool
def search_wikipedia(query: str) -> str:
    """Search Wikipedia for a given query. Returns a summary."""
    import urllib.request
    import urllib.parse

    try:
        encoded = urllib.parse.quote(query)
        url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": "LangChainAI/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            if "extract" in data:
                return data["extract"][:1500]
            elif "title" in data:
                return f"Topik: {data['title']}. Tidak ada ringkasan tersedia."
            else:
                return "Tidak ditemukan."
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return f"Tidak ditemukan artikel Wikipedia untuk '{query}'."
        return f"Error Wikipedia: {e}"
    except Exception as e:
        return f"Error: {e}"


@tool
def get_weather(city: str) -> str:
    """Get current weather for a city. Input: city name."""
    import urllib.request
    import urllib.parse

    try:
        encoded = urllib.parse.quote(city)
        url = f"https://wttr.in/{encoded}?format=%C+|+%t+|+%h+humidity+|+%w+wind"
        req = urllib.request.Request(url, headers={"User-Agent": "curl/8.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = resp.read().decode().strip()
            return f"Cuaca di {city}: {result}"
    except Exception as e:
        return f"Error getting weather: {e}"


TOOLS = [calculate, get_current_time, search_wikipedia, get_weather]


# ═══════════════════════════════════════════════════════════════
#  Conversation Memory (Session-based)
# ═══════════════════════════════════════════════════════════════

from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

# In-memory session store: {session_id: [messages]}
sessions_store: dict[str, list] = {}


def get_session_history(session_id: str) -> list:
    """Get or create a session's message history."""
    if session_id not in sessions_store:
        sessions_store[session_id] = []
    return sessions_store[session_id]


def add_to_session(session_id: str, role: str, content: str):
    """Add a message to session history, trimming oldest if needed."""
    history = get_session_history(session_id)
    history.append({"role": role, "content": content})
    # Keep only the last N messages
    if len(history) > MAX_HISTORY_PER_SESSION:
        sessions_store[session_id] = history[-MAX_HISTORY_PER_SESSION:]


def build_messages_from_history(session_id: str, system_prompt: str) -> list:
    """Build a list of LangChain message objects from session history + system prompt."""
    history = get_session_history(session_id)
    messages = [SystemMessage(content=system_prompt)]

    for msg in history[-MAX_HISTORY_PER_SESSION:]:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        else:
            messages.append(AIMessage(content=msg["content"]))

    return messages


def list_sessions() -> list[dict]:
    """List all active sessions with metadata."""
    result = []
    for sid, msgs in sessions_store.items():
        first_user_msg = ""
        for m in msgs:
            if m["role"] == "user":
                first_user_msg = m["content"][:50]
                break
        result.append({
            "session_id": sid,
            "message_count": len(msgs),
            "preview": first_user_msg,
        })
    return result


# ═══════════════════════════════════════════════════════════════
#  RAG (in-memory vector store)
# ═══════════════════════════════════════════════════════════════

from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# Global RAG state
rag_vectorstore = None
rag_retriever = None
rag_chain = None
rag_current_provider = None
rag_embeddings = None


def init_rag():
    """Initialize RAG with default documents."""
    global rag_retriever, rag_embeddings

    documents = [
        Document(
            page_content="LangChain adalah framework open-source untuk membangun aplikasi "
                         "berbasis LLM. Fitur utama: chains, agents, RAG, memory, tool calling.",
            metadata={"source": "docs", "topic": "introduction"},
        ),
        Document(
            page_content="RAG (Retrieval Augmented Generation) adalah pola arsitektur yang "
                         "menggabungkan information retrieval dengan generative LLM. "
                         "Cocok untuk QA atas dokumen privat tanpa fine-tuning.",
            metadata={"source": "docs", "topic": "rag"},
        ),
        Document(
            page_content="LangGraph adalah library untuk membangun agent multi-langkah "
                         "dengan state graph. Mendukung cycles, branching, dan persistence.",
            metadata={"source": "docs", "topic": "langgraph"},
        ),
        Document(
            page_content="FastAPI adalah framework Python modern untuk membangun API REST. "
                         "Cepat (high performance), mudah digunakan, dengan dokumentasi "
                         "OpenAPI otomatis dan validasi request via Pydantic.",
            metadata={"source": "docs", "topic": "fastapi"},
        ),
    ]

    rag_embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=20)
    chunks = text_splitter.split_documents(documents)
    rag_vectorstore = FAISS.from_documents(chunks, rag_embeddings)
    rag_retriever = rag_vectorstore.as_retriever(search_kwargs={"k": 3})

    log.info("RAG initialized with %d chunks", len(chunks))


def format_docs(docs):
    """Format retrieved documents into a single context string."""
    return "\n\n".join(d.page_content for d in docs)


def rebuild_rag_chain(provider: str = "ollama", **llm_kwargs):
    """Rebuild the RAG QA chain (call after ingesting new docs or changing provider)."""
    global rag_chain, rag_current_provider

    if rag_retriever is None:
        init_rag()

    llm = create_llm(provider=provider, **llm_kwargs)

    prompt = ChatPromptTemplate.from_messages([
        ("system", "Jawab pertanyaan berdasarkan konteks berikut:\n\n{context}"),
        ("human", "{input}"),
    ])

    rag_chain = (
        {"context": rag_retriever | format_docs, "input": RunnablePassthrough()}
        | prompt
        | llm
    )
    rag_current_provider = provider
    log.info("RAG chain rebuilt with provider=%s, kwargs=%s", provider, llm_kwargs)


def ingest_text(text: str, source: str = "upload", provider: str = "ollama") -> dict:
    """Ingest raw text into RAG vector store."""
    global rag_vectorstore, rag_retriever, rag_embeddings

    if rag_embeddings is None:
        rag_embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

    doc = Document(page_content=text, metadata={"source": source})
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_documents([doc])

    if rag_vectorstore is None:
        rag_vectorstore = FAISS.from_documents(chunks, rag_embeddings)
    else:
        rag_vectorstore.add_documents(chunks)

    rag_retriever = rag_vectorstore.as_retriever(search_kwargs={"k": 3})
    rebuild_rag_chain(provider=provider, temperature=0.0)

    return {"ingested": 1, "chunks": len(chunks)}


# ═══════════════════════════════════════════════════════════════
#  Application Lifecycle
# ═══════════════════════════════════════════════════════════════

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup: init RAG on boot."""
    log.info("Starting LangChain FastAPI Server...")
    init_rag()
    rebuild_rag_chain(provider="ollama", temperature=0.0)
    yield
    log.info("Shutting down.")


# ── Detect available providers ───────────────────────────────

def check_ollama_available() -> bool:
    """Check if Ollama is reachable at the configured URL."""
    import urllib.request
    base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    try:
        req = urllib.request.Request(f"{base_url}/api/tags", method="GET")
        with urllib.request.urlopen(req, timeout=3) as resp:
            return resp.status == 200
    except Exception:
        return False


def detect_fallback_providers() -> dict:
    """Detect which LLM providers are configured and available.
    Returns a dict with keys: preferred, available."""
    providers = []

    # Check Ollama (local or cloud)
    ollama_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    if check_ollama_available():
        providers.append("ollama")
        log.info("✅ Ollama detected at %s", ollama_url)
    else:
        log.info("❌ Ollama not reachable at %s", ollama_url)

    # Check other providers by API key presence
    if os.getenv("OPENAI_API_KEY"):
        providers.append("openai")
        log.info("✅ OpenAI API key found")
    if os.getenv("ANTHROPIC_API_KEY"):
        providers.append("anthropic")
        log.info("✅ Anthropic API key found")
    if os.getenv("GOOGLE_API_KEY"):
        providers.append("google")
        log.info("✅ Google API key found")

    # Default to ollama even if not detected (will fail gracefully)
    if not providers:
        providers.append("ollama")
        log.warning("⚠️  No provider detected! Will use Ollama (might fail on cloud)")
        log.warning("   Set OPENAI_API_KEY, ANTHROPIC_API_KEY, or GOOGLE_API_KEY in env")
        log.warning("   Or route Ollama to your cloud instance via OLLAMA_BASE_URL")

    return {"preferred": providers[0], "available": providers}


app = FastAPI(
    title="LangChain AI API",
    description="Full-stack AI server: Chat, Agent with tools, RAG QA, PDF upload.",
    version="2.0.0",
    lifespan=lifespan,
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (frontend)
STATIC_DIR = Path(__file__).parent / "static"
if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
    log.info("Static files mounted from %s", STATIC_DIR)

# Detect deployment environment and available providers
IS_HF_SPACE = bool(os.getenv("SPACE_ID"))
if IS_HF_SPACE:
    log.info("🌐 Running on Hugging Face Spaces!")

AVAILABLE_PROVIDERS = detect_fallback_providers()
log.info("📡 Preferred provider: %s | Available: %s",
         AVAILABLE_PROVIDERS["preferred"], AVAILABLE_PROVIDERS["available"])


# ═══════════════════════════════════════════════════════════════
#  Pydantic Schemas
# ═══════════════════════════════════════════════════════════════

class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, description="User message")
    session_id: str = Field(default_factory=lambda: str(uuid.uuid4())[:8], description="Session ID for conversation memory")
    provider: str = Field("ollama", description="LLM provider")
    system_prompt: str = Field(
        "Kamu adalah asisten AI yang ramah dan membantu. Jawab dalam Bahasa Indonesia.",
        description="System prompt override",
    )
    temperature: float = Field(0.7, ge=0.0, le=2.0)
    model: str | None = Field(None, description="Model name override")

    model_config = {"json_schema_extra": {
        "example": {
            "message": "Apa itu LangChain?",
            "session_id": "abc12345",
            "provider": "ollama",
        }
    }}


class AgentRequest(BaseModel):
    message: str = Field(..., min_length=1, description="User message")
    session_id: str = Field(default_factory=lambda: str(uuid.uuid4())[:8], description="Session ID")
    provider: str = Field("ollama", description="LLM provider")
    temperature: float = Field(0.0, ge=0.0, le=2.0)
    model: str | None = Field(None)

    model_config = {"json_schema_extra": {
        "example": {
            "message": "Hitung 2 + 3 * 4 dan cari Wikipedia tentang AI",
            "session_id": "abc12345",
            "provider": "ollama",
        }
    }}


class RagQueryRequest(BaseModel):
    query: str = Field(..., min_length=1, description="Question to ask")
    provider: str = Field("ollama", description="LLM provider")
    temperature: float = Field(0.0, ge=0.0, le=2.0)
    model: str | None = Field(None)

    model_config = {"json_schema_extra": {
        "example": {"query": "Apa itu RAG?"}
    }}


class DocumentSchema(BaseModel):
    page_content: str = Field(..., min_length=1, description="Document text content")
    metadata: dict = Field(default_factory=dict, description="Optional metadata")


class RagIngestRequest(BaseModel):
    documents: list[DocumentSchema] = Field(..., min_length=1, description="Documents to ingest")
    provider: str = Field("ollama")


# Response models
class ChatResponse(BaseModel):
    response: str
    session_id: str
    provider: str
    model: str


class AgentResponse(BaseModel):
    response: str
    session_id: str
    tool_calls: list[dict] = []
    provider: str
    model: str


class RagResponse(BaseModel):
    answer: str
    sources: list[dict] = []
    provider: str


# ═══════════════════════════════════════════════════════════════
#  Endpoints — Frontend
# ═══════════════════════════════════════════════════════════════

@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    """Serve the chat UI or API info."""
    accept = request.headers.get("accept", "")
    index_path = STATIC_DIR / "index.html"

    # If browser requests HTML, serve the frontend (cached after first read)
    if "text/html" in accept and index_path.exists():
        if not hasattr(root, "_html_cache") or root._html_cache is None:
            root._html_cache = index_path.read_text(encoding="utf-8")
        return HTMLResponse(content=root._html_cache)

    # Otherwise return API info
    return HTMLResponse(
        content=json.dumps({
            "app": "LangChain AI API",
            "version": "2.0.0",
            "docs": "/docs",
            "frontend": "/ (buka di browser)",
            "endpoints": {
                "GET  /health": "Server health check",
                "POST /chat": "Chat with conversation memory",
                "POST /agent": "Agent with tools (calculator, weather, wikipedia, time)",
                "GET  /agent/stream": "SSE streaming agent",
                "POST /rag/query": "RAG QA over documents",
                "POST /rag/ingest": "Ingest text documents into RAG",
                "POST /rag/upload": "Upload PDF/DOCX/TXT file into RAG",
                "GET  /sessions": "List active conversation sessions",
            }
        }, indent=2),
        media_type="application/json",
    )


# ═══════════════════════════════════════════════════════════════
#  Endpoints — Health & Sessions
# ═══════════════════════════════════════════════════════════════

@app.get("/health")
async def health():
    """Health check endpoint."""
    return {
        "status": "ok",
        "version": "2.0.0",
        "rag_ready": rag_chain is not None,
        "sessions_active": len(sessions_store),
        "provider": AVAILABLE_PROVIDERS["preferred"],
        "providers_available": AVAILABLE_PROVIDERS["available"],
        "is_hf_space": IS_HF_SPACE,
    }


@app.get("/sessions")
async def sessions():
    """List all active conversation sessions."""
    return {"sessions": list_sessions(), "total": len(sessions_store)}


@app.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    """Delete a conversation session."""
    if session_id in sessions_store:
        del sessions_store[session_id]
        return {"status": "deleted", "session_id": session_id}
    raise HTTPException(status_code=404, detail="Session not found")


# ═══════════════════════════════════════════════════════════════
#  Endpoints — Chat (with memory)
# ═══════════════════════════════════════════════════════════════

@app.post("/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    """Chat with LLM, with conversation memory."""
    try:
        llm = create_llm(
            provider=req.provider,
            temperature=req.temperature,
            model=req.model,
        )

        # Build messages from session history
        messages = build_messages_from_history(req.session_id, req.system_prompt)
        messages.append(HumanMessage(content=req.message))

        result = llm.invoke(messages)
        model_used = result.response_metadata.get("model", req.model or "")

        # Save to session history
        add_to_session(req.session_id, "user", req.message)
        add_to_session(req.session_id, "assistant", result.content)

        return ChatResponse(
            response=result.content,
            session_id=req.session_id,
            provider=req.provider,
            model=model_used or req.provider,
        )

    except Exception as e:
        error_msg = str(e)
        log.exception("Chat error")
        if "subscription" in error_msg.lower() or "403" in error_msg:
            raise HTTPException(
                status_code=402,
                detail=f"Model requires subscription. Gunakan model gratis: qwen3-coder-next:cloud",
            )
        if "connect" in error_msg.lower() or "refused" in error_msg.lower():
            raise HTTPException(
                status_code=503,
                detail=f"Tidak bisa connect ke Ollama. Apakah Ollama sudah jalan?",
            )
        raise HTTPException(status_code=500, detail=error_msg[:500])


# ═══════════════════════════════════════════════════════════════
#  Endpoints — Agent (with tools + memory)
# ═══════════════════════════════════════════════════════════════

@app.post("/agent", response_model=AgentResponse)
async def agent(req: AgentRequest):
    """Agent with tools (calculator, weather, wikipedia, time)."""
    try:
        llm = create_llm(
            provider=req.provider,
            temperature=req.temperature,
            model=req.model,
        )

        from langgraph.prebuilt import create_react_agent

        # Build conversation history
        history = get_session_history(req.session_id)
        langchain_messages = []
        for msg in history[-10:]:  # Use last 10 for agent context
            if msg["role"] == "user":
                langchain_messages.append(HumanMessage(content=msg["content"]))
            else:
                langchain_messages.append(AIMessage(content=msg["content"]))
        langchain_messages.append(HumanMessage(content=req.message))

        agent_executor = create_react_agent(llm, TOOLS)

        result = await agent_executor.ainvoke(
            {"messages": langchain_messages}
        )

        final_msg = result["messages"][-1]
        tool_calls = []

        for msg in result["messages"]:
            if hasattr(msg, "tool_calls") and msg.tool_calls:
                for tc in msg.tool_calls:
                    tool_calls.append({
                        "name": tc.get("name", ""),
                        "args": tc.get("args", {}),
                    })

        # Save to session
        add_to_session(req.session_id, "user", req.message)
        add_to_session(req.session_id, "assistant", final_msg.content)

        return AgentResponse(
            response=final_msg.content,
            session_id=req.session_id,
            tool_calls=tool_calls,
            provider=req.provider,
            model=req.model or req.provider,
        )

    except Exception as e:
        error_msg = str(e)
        log.exception("Agent error")
        if "subscription" in error_msg.lower() or "403" in error_msg:
            raise HTTPException(
                status_code=402,
                detail="Model requires subscription. Gunakan model gratis: qwen3-coder-next:cloud",
            )
        if "connect" in error_msg.lower() or "refused" in error_msg.lower():
            raise HTTPException(
                status_code=503,
                detail="Tidak bisa connect ke Ollama. Apakah Ollama sudah jalan?",
            )
        raise HTTPException(status_code=500, detail=error_msg[:500])


@app.get("/agent/stream")
async def agent_stream(message: str, session_id: str = "default", provider: str = "ollama"):
    """SSE streaming agent response."""
    try:
        llm = create_llm(provider=provider, temperature=0.0)

        from langgraph.prebuilt import create_react_agent
        from langchain_core.messages import HumanMessage

        agent_executor = create_react_agent(llm, TOOLS)

        async def event_generator():
            yield {"event": "start", "data": json.dumps({"status": "processing"})}

            async for chunk in agent_executor.astream(
                {"messages": [HumanMessage(content=message)]}
            ):
                for node_name, node_data in chunk.items():
                    if "messages" in node_data:
                        for msg in node_data["messages"]:
                            content = msg.content if hasattr(msg, "content") else str(msg)
                            msg_type = msg.type if hasattr(msg, "type") else node_name

                            if content:
                                yield {
                                    "event": "message",
                                    "data": json.dumps({
                                        "role": msg_type,
                                        "content": content,
                                        "node": node_name,
                                    }),
                                }

            # Save to session
            add_to_session(session_id, "user", message)

            yield {"event": "done", "data": json.dumps({"status": "completed"})}

        return EventSourceResponse(event_generator())

    except Exception as e:
        log.exception("Stream error")
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
#  Endpoints — RAG
# ═══════════════════════════════════════════════════════════════

@app.post("/rag/query", response_model=RagResponse)
async def rag_query(req: RagQueryRequest):
    """Query the RAG system."""
    global rag_chain, rag_current_provider

    try:
        if rag_chain is None or rag_current_provider != req.provider:
            rebuild_rag_chain(provider=req.provider, temperature=req.temperature, model=req.model)

        retrieved_docs = await rag_retriever.ainvoke(req.query)
        sources = [
            {
                "content": doc.page_content[:200],
                "metadata": doc.metadata,
            }
            for doc in retrieved_docs
        ]

        result = await rag_chain.ainvoke(req.query)

        return RagResponse(
            answer=result.content if hasattr(result, "content") else str(result),
            sources=sources,
            provider=req.provider,
        )

    except Exception as e:
        error_msg = str(e)
        log.exception("RAG query error")
        if "subscription" in error_msg.lower() or "403" in error_msg:
            raise HTTPException(status_code=402, detail="Model requires subscription")
        if "connect" in error_msg.lower() or "refused" in error_msg.lower():
            raise HTTPException(status_code=503, detail="Cannot connect to Ollama")
        raise HTTPException(status_code=500, detail=error_msg[:500])


@app.post("/rag/ingest")
async def rag_ingest(req: RagIngestRequest):
    """Ingest text documents into RAG vector store."""
    global rag_vectorstore, rag_retriever, rag_embeddings

    try:
        documents = [
            Document(page_content=d.page_content, metadata=d.metadata)
            for d in req.documents
        ]

        if rag_embeddings is None:
            rag_embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_documents(documents)

        if rag_vectorstore is None:
            rag_vectorstore = FAISS.from_documents(chunks, rag_embeddings)
        else:
            rag_vectorstore.add_documents(chunks)
            log.info("Merged %d chunks into existing vector store", len(chunks))

        rag_retriever = rag_vectorstore.as_retriever(search_kwargs={"k": 3})
        rebuild_rag_chain(provider=req.provider, temperature=0.0)

        return {
            "status": "ok",
            "ingested": len(documents),
            "chunks": len(chunks),
            "message": f"Successfully ingested {len(documents)} documents ({len(chunks)} chunks)",
        }

    except Exception as e:
        log.exception("RAG ingest error")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/rag/upload")
async def rag_upload(file: UploadFile = File(...), provider: str = "ollama"):
    """Upload PDF/DOCX/TXT file and ingest into RAG."""
    try:
        # Validate file type
        ext = Path(file.filename).suffix.lower() if file.filename else ".txt"

        # Save uploaded file temporarily
        temp_path = UPLOAD_DIR / f"{uuid.uuid4()}{ext}"
        content_bytes = await file.read()

        with open(temp_path, "wb") as f:
            f.write(content_bytes)

        # Extract text based on file type
        text = ""

        if ext == ".txt":
            text = content_bytes.decode("utf-8", errors="replace")

        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(temp_path)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Gagal membaca PDF: {e}. Pastikan file PDF valid.",
                )

        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(temp_path))
                text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="python-docx tidak terinstall. Install: pip install python-docx",
                )
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Gagal membaca DOCX: {e}",
                )
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Format file tidak didukung: {ext}. Gunakan .pdf, .docx, atau .txt",
            )

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="Tidak ada teks yang bisa diekstrak dari file tersebut.",
            )

        # Ingest into RAG
        result = ingest_text(text, source=file.filename or "upload", provider=provider)

        # Clean up temp file
        temp_path.unlink(missing_ok=True)

        return {
            "status": "ok",
            "filename": file.filename,
            "characters": len(text),
            **result,
            "message": f"✅ Berhasil meng-upload '{file.filename}' ({len(text)} karakter, {result['chunks']} chunks)",
        }

    except HTTPException:
        raise
    except Exception as e:
        log.exception("RAG upload error")
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "server:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info",
    )
